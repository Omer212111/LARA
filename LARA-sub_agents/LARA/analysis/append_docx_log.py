"""
LARA — Append an Experiment Log entry to "LARA - Documentation.docx"

Follows the document's existing five-section template:

    Experiment Log: <name>            (Heading 1)
    Date: ... | Conducted by: ...
    1. Objective & Hypothesis
    2. Environment Setup               (Heading 3)
    3. Quantitative Results
    4. Qualitative Analysis            (Heading 3)
    5. Conclusions & Next Steps        (Heading 3)

Content is supplied as a JSON payload so the writer stays generic; see
`analysis/docx_entries/` for the entries used in the 2026-07-20 study.

CLI:
    python analysis/append_docx_log.py analysis/docx_entries/spotify.json
    python analysis/append_docx_log.py entry.json --docx "LARA - Documentation.docx" --dry-run
"""
from __future__ import annotations

import argparse
import json
import shutil
from datetime import datetime
from pathlib import Path

import docx  # python-docx

HERE = Path(__file__).parent
LARA_ROOT = HERE.parent
DEFAULT_DOCX = LARA_ROOT / "LARA - Documentation.docx"


def _add_bullets(document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(f"• {item}")


def _add_table(document, table_spec: dict) -> None:
    """table_spec = {"headers": [...], "rows": [[...], ...]}"""
    headers = table_spec["headers"]
    rows = table_spec["rows"]
    table = document.add_table(rows=1, cols=len(headers))
    # "TableNormal" is the only table style this document defines, and it is what
    # every existing experiment-log table uses. "Table Grid" is absent and raises.
    try:
        table.style = "TableNormal"
    except KeyError:
        pass  # fall back to the default style rather than fail the whole append
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = str(text)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = str(text)


def _add_block(document, block) -> None:
    """A block is a plain string, or {"bullets": [...]} / {"table": {...}}."""
    if isinstance(block, str):
        document.add_paragraph(block)
    elif "bullets" in block:
        _add_bullets(document, block["bullets"])
    elif "table" in block:
        _add_table(document, block["table"])
    else:
        raise ValueError(f"Unrecognised block: {block!r}")


def append_entry(entry: dict, docx_path: Path, dry_run: bool = False) -> None:
    if dry_run:
        print(f"[dry-run] Would append 'Experiment Log: {entry['name']}' to {docx_path}")
        for section in ("objective", "environment", "quantitative",
                        "qualitative", "conclusions"):
            blocks = entry.get(section, [])
            print(f"  {section}: {len(blocks)} block(s)")
        return

    # The document is the project's record of every experiment — never write it
    # in place without a timestamped backup first.
    backup = docx_path.with_suffix(
        f".backup-{datetime.now():%Y%m%d-%H%M%S}.docx")
    shutil.copy2(docx_path, backup)

    document = docx.Document(str(docx_path))

    document.add_heading(f"Experiment Log: {entry['name']}", level=1)
    document.add_paragraph(
        f"Date: {entry['date']} | Conducted by: {entry.get('author', 'Omer')}"
    )

    document.add_paragraph("1. Objective & Hypothesis")
    for block in entry.get("objective", []):
        _add_block(document, block)

    document.add_heading("2. Environment Setup", level=3)
    for block in entry.get("environment", []):
        _add_block(document, block)

    document.add_paragraph("3. Quantitative Results")
    for block in entry.get("quantitative", []):
        _add_block(document, block)

    document.add_heading("4. Qualitative Analysis", level=3)
    for block in entry.get("qualitative", []):
        _add_block(document, block)

    document.add_heading("5. Conclusions & Next Steps", level=3)
    for block in entry.get("conclusions", []):
        _add_block(document, block)

    document.save(str(docx_path))
    print(f"Appended 'Experiment Log: {entry['name']}' to {docx_path}")
    print(f"  Backup: {backup}")


def _cli() -> None:
    p = argparse.ArgumentParser(description=__doc__,
                                formatter_class=argparse.RawDescriptionHelpFormatter)
    p.add_argument("entry", type=Path, help="JSON file describing the entry")
    p.add_argument("--docx", type=Path, default=DEFAULT_DOCX)
    p.add_argument("--dry-run", action="store_true")
    args = p.parse_args()

    entry = json.loads(args.entry.read_text(encoding="utf-8"))
    append_entry(entry, args.docx, dry_run=args.dry_run)


if __name__ == "__main__":
    _cli()
