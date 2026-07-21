"""
LARA — Rewrite sections 4 & 5 of an Experiment Log entry in place

The append-only writer (append_docx_log.py) cannot revise an entry once written.
This script replaces the body paragraphs under "4. Qualitative Analysis" and
"5. Conclusions & Next Steps" for a named entry, leaving every other section,
table and entry untouched.

Body paragraphs are deleted at the XML level (python-docx has no delete API) and
new ones inserted before the following heading, so ordering is preserved.

CLI:
    python analysis/rewrite_docx_sections.py content.json [--dry-run]

content.json:
    {"entries": [{"match": "<substring of the entry heading>",
                  "qualitative": ["...", ...],
                  "conclusions": ["...", ...]}]}
"""
from __future__ import annotations

import argparse
import copy
import json
import shutil
from datetime import datetime
from pathlib import Path

import docx

HERE = Path(__file__).parent
DEFAULT_DOCX = HERE.parent / "LARA - Documentation.docx"

SECTION_4 = "4. Qualitative Analysis"
SECTION_5 = "5. Conclusions & Next Steps"


def _paragraph_index(paragraphs, predicate, start: int = 0) -> int | None:
    for i in range(start, len(paragraphs)):
        if predicate(paragraphs[i]):
            return i
    return None


def _delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def _insert_after(anchor_paragraph, text: str):
    """Insert a new paragraph directly after `anchor_paragraph`, same style."""
    new_p = copy.deepcopy(anchor_paragraph._element)
    # Strip all runs from the clone, then set our text.
    for child in list(new_p):
        if child.tag.endswith('}r'):
            new_p.remove(child)
    anchor_paragraph._element.addnext(new_p)

    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, anchor_paragraph._parent)
    para.text = text
    para.style = anchor_paragraph.style
    return para


def rewrite_section(document, heading_text: str, entry_start: int,
                    new_body: list[str], stop_at: str | None) -> int:
    """Replace the body under `heading_text`. Returns paragraphs replaced."""
    paragraphs = document.paragraphs
    head_i = _paragraph_index(
        paragraphs, lambda p: p.text.strip() == heading_text, entry_start)
    if head_i is None:
        raise SystemExit(f"Section not found after index {entry_start}: {heading_text}")

    # Body runs until the next section heading, the next entry, or the document end.
    end_i = len(paragraphs)
    for i in range(head_i + 1, len(paragraphs)):
        text = paragraphs[i].text.strip()
        if (stop_at and text == stop_at) or text.startswith("Experiment Log:"):
            end_i = i
            break

    removed = 0
    for para in paragraphs[head_i + 1:end_i]:
        _delete_paragraph(para)
        removed += 1

    anchor = document.paragraphs[head_i]
    for text in new_body:
        anchor = _insert_after(anchor, text)
    return removed


def apply(content: dict, docx_path: Path, dry_run: bool = False) -> None:
    document = docx.Document(str(docx_path))

    for entry in content["entries"]:
        match = entry["match"]
        start = _paragraph_index(
            document.paragraphs,
            lambda p: p.text.startswith("Experiment Log:") and match in p.text)
        if start is None:
            raise SystemExit(f"Entry not found: {match!r}")

        if dry_run:
            print(f"[dry-run] {match}: "
                  f"section 4 -> {len(entry['qualitative'])} paras, "
                  f"section 5 -> {len(entry['conclusions'])} paras")
            continue

        # Section 5 first: rewriting it does not shift section 4's index.
        n5 = rewrite_section(document, SECTION_5, start, entry["conclusions"], None)
        n4 = rewrite_section(document, SECTION_4, start, entry["qualitative"], SECTION_5)
        print(f"{match}: replaced {n4} paras in section 4, {n5} in section 5")

    if dry_run:
        return

    backup = docx_path.with_suffix(f".backup-{datetime.now():%Y%m%d-%H%M%S}.docx")
    shutil.copy2(docx_path, backup)
    document.save(str(docx_path))
    print(f"Saved {docx_path}\n  Backup: {backup}")


def _cli() -> None:
    p = argparse.ArgumentParser(description=__doc__,
                                formatter_class=argparse.RawDescriptionHelpFormatter)
    p.add_argument("content", type=Path)
    p.add_argument("--docx", type=Path, default=DEFAULT_DOCX)
    p.add_argument("--dry-run", action="store_true")
    args = p.parse_args()
    apply(json.loads(args.content.read_text(encoding="utf-8")), args.docx, args.dry_run)


if __name__ == "__main__":
    _cli()
